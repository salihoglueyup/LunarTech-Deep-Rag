"""
LunarTech AI — Document Processor v3
Multi-format: PDF, DOCX, TXT, MD support.
"""

import os
import re
import pdfplumber
from io import BytesIO
from utils.helpers import clean_text, chunk_text

# ── Supported formats ──
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}


def get_file_extension(filename: str) -> str:
    return os.path.splitext(filename.lower())[1]


def is_supported(filename: str) -> bool:
    return get_file_extension(filename) in SUPPORTED_EXTENSIONS


# ══════════════════════════════════════════════════════════
# PDF PROCESSING
# ══════════════════════════════════════════════════════════


def _extract_pdf(file_obj) -> dict:
    """Extracts text, tables, and metadata from PDF."""
    pages = []
    full_text_parts = []
    metadata = {}
    tables_found = []

    if isinstance(file_obj, str):
        pdf = pdfplumber.open(file_obj)
    else:
        pdf = pdfplumber.open(BytesIO(file_obj.read()))

    try:
        if pdf.metadata:
            metadata = {
                "title": pdf.metadata.get("Title", ""),
                "author": pdf.metadata.get("Author", ""),
                "creator": pdf.metadata.get("Creator", ""),
                "subject": pdf.metadata.get("Subject", ""),
            }

        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            page_text = clean_text(page_text)

            # OCR Fallback for image-based PDFs
            if not page_text.strip():
                try:
                    import pytesseract

                    # pdfplumber to_image() creates a PageImage object. .original returns the PIL image.
                    pil_img = page.to_image(resolution=200).original
                    page_text = pytesseract.image_to_string(pil_img, lang="tur+eng")
                    page_text = clean_text(page_text)
                except Exception:
                    pass

            # Table extraction (native pdfplumber)
            page_tables = page.extract_tables()
            if page_tables:
                for tbl in page_tables:
                    tables_found.append(
                        {
                            "page": i + 1,
                            "rows": len(tbl),
                            "data": tbl[:5],  # first 5 rows
                        }
                    )

            if page_text:
                pages.append(
                    {
                        "page_num": i + 1,
                        "text": page_text,
                        "char_count": len(page_text),
                        "has_tables": len(page_tables) > 0 if page_tables else False,
                    }
                )
                full_text_parts.append(f"[Page {i+1}]\n{page_text}")
    finally:
        pdf.close()

    return {
        "full_text": "\n\n".join(full_text_parts),
        "pages": pages,
        "page_count": len(pages),
        "metadata": metadata,
        "tables": tables_found,
        "format": "pdf",
    }


# ══════════════════════════════════════════════════════════
# DOCX PROCESSING
# ══════════════════════════════════════════════════════════


def _extract_docx(file_obj) -> dict:
    """Extracts text from DOCX file."""
    try:
        from docx import Document
    except ImportError:
        return {
            "error": "python-docx library is not installed. pip install python-docx"
        }

    if isinstance(file_obj, str):
        doc = Document(file_obj)
    else:
        doc = Document(BytesIO(file_obj.read()))

    full_text_parts = []
    headings = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Heading detection
        if para.style and para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading", "").strip()
            level = int(level) if level.isdigit() else 1
            headings.append({"level": level, "text": text})
            full_text_parts.append(f"{'#' * level} {text}")
        else:
            full_text_parts.append(text)

    # Table extraction
    tables_found = []
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows[:5]:
            rows.append([cell.text.strip() for cell in row.cells])
        tables_found.append({"index": i + 1, "rows": len(table.rows), "data": rows})

    full_text = "\n\n".join(full_text_parts)

    # Metadata
    metadata = {}
    if doc.core_properties:
        metadata = {
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "subject": doc.core_properties.subject or "",
        }

    return {
        "full_text": full_text,
        "pages": [{"page_num": 1, "text": full_text, "char_count": len(full_text)}],
        "page_count": 1,
        "metadata": metadata,
        "headings": headings,
        "tables": tables_found,
        "format": "docx",
    }


# ══════════════════════════════════════════════════════════
# TXT / MARKDOWN PROCESSING
# ══════════════════════════════════════════════════════════


def _extract_text_file(file_obj, ext: str) -> dict:
    """Extracts text from TXT or Markdown file."""
    if isinstance(file_obj, str):
        with open(file_obj, "r", encoding="utf-8") as f:
            content = f.read()
    else:
        content = file_obj.read().decode("utf-8")

    content = clean_text(content)

    # Heading detection (markdown)
    headings = []
    if ext in (".md", ".markdown"):
        for line in content.split("\n"):
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                text = line.lstrip("# ").strip()
                if text:
                    headings.append({"level": level, "text": text})

    return {
        "full_text": content,
        "pages": [{"page_num": 1, "text": content, "char_count": len(content)}],
        "page_count": 1,
        "metadata": {},
        "headings": headings,
        "tables": [],
        "format": ext.lstrip("."),
    }


# ══════════════════════════════════════════════════════════
# UNIFIED API
# ══════════════════════════════════════════════════════════


def extract_text(file_obj, filename: str = None) -> dict:
    """
    Automatic text extraction based on file format.

    Returns:
        {
            "full_text": str,
            "pages": list,
            "page_count": int,
            "metadata": dict,
            "format": str,
            "tables": list,
            "headings": list,
        }
    """
    if filename is None:
        filename = getattr(file_obj, "name", "unknown.txt")

    ext = get_file_extension(filename)

    if ext == ".pdf":
        return _extract_pdf(file_obj)
    elif ext == ".docx":
        return _extract_docx(file_obj)
    elif ext in (".txt", ".md", ".markdown"):
        return _extract_text_file(file_obj, ext)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def process_to_chunks(
    file_obj, filename: str = None, chunk_size: int = 1000, overlap: int = 200
) -> dict:
    """Extracts the file and splits it into chunks."""
    result = extract_text(file_obj, filename)
    chunks = chunk_text(result["full_text"], chunk_size=chunk_size, overlap=overlap)

    return {
        **result,
        "chunks": chunks,
        "chunk_count": len(chunks),
    }


# ── Quality Score ──


def document_quality_score(result: dict) -> dict:
    """Calculates document quality score."""
    text = result.get("full_text", "")
    score = 100
    issues = []

    # Text length
    if len(text) < 100:
        score -= 30
        issues.append("Very short text")
    elif len(text) < 500:
        score -= 15
        issues.append("Short text")

    # Empty page ratio
    pages = result.get("pages", [])
    if pages:
        empty = sum(1 for p in pages if len(p.get("text", "")) < 20)
        empty_ratio = empty / len(pages)
        if empty_ratio > 0.3:
            score -= 20
            issues.append(f"High empty page ratio ({int(empty_ratio*100)}%)")

    # Repetition check
    words = text.split()
    if len(words) > 50:
        unique_ratio = len(set(words)) / len(words)
        if unique_ratio < 0.3:
            score -= 15
            issues.append("Too much repetition")

    # Table presence (bonus)
    if result.get("tables"):
        score = min(100, score + 5)

    return {
        "score": max(0, min(100, score)),
        "grade": (
            "A" if score >= 90 else "B" if score >= 70 else "C" if score >= 50 else "D"
        ),
        "issues": issues,
        "format": result.get("format", "unknown"),
    }
